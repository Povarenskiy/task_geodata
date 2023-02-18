import time
import json
import folium
from rosreestr2coord import Area
from docx import Document
from docx.shared import Inches, Pt
from selenium import webdriver


def _get_map(cadastral_number):
    # получение географических данных в формате json c rosreestr2coord
    area = Area(cadastral_number).to_geojson_poly() 
    if area:
        data = json.loads(area)
        
        # получение центра области для отображения карты
        x = data['properties']['center'].get('x')
        y = data['properties']['center'].get('y')

        # получение карты для отображения на html странице
        m = folium.Map(location=[y , x], zoom_start=17)
        folium.GeoJson(data, name="geojson").add_to(m)
        return m


def get_map(cadastral_number):
    start = time.time()
    m = _get_map(cadastral_number)

    # запись времени выоплнения
    result = time.time() - start
    result = '{:.2e}'.format(result) if result < 1 else round(result, 2)
    
    if m:
        return {'map_html': m._repr_html_(), 'result': result}
    else:
        return {'error': 'Не удалось загрузить карту по кадастровому номеру'}

    
def create_docx(cadastral_number=None, title=None, text=None):
    document = Document()
    
    document.styles['Normal'].font.size = Pt(14)
    document.styles['Title'].font.size = Pt(18)

    if title:
        document.add_heading(title.capitalize(), 0)

    if cadastral_number:
        p = document.add_paragraph()
        p.add_run('Кадастровый номер: ').bold = True
        p.add_run(cadastral_number)

    if text:
        p = document.add_paragraph()
        p.add_run('Описание: ').bold = True
        p.add_run(text)
    
    m = _get_map(cadastral_number)
    if m:
        options = webdriver.chrome.options.Options()
        options.add_argument("--headless")
        driver = webdriver.Chrome(options=options)

        img = m._to_png(delay=3, driver=driver)

        with open('map.png', 'wb') as map_png:
            map_png.write(img)
            document.add_picture('map.png', width=Inches(5.4), height=Inches(3))

    path = 'demo.docx'
    document.save(path)
    return path
